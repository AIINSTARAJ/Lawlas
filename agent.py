#!/usr/bin/env python3
"""
LawLas ⚖️ - AI-Powered Legal Assistant
A comprehensive legal assistant system using LangChain and Gemini models.
"""

import os
import json
import hashlib
import time
import requests
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

# Install required packages
def install_requirements():
    packages = [
        'langchain', 'langchain-google-genai', 'langchain-community',
        'colored', 'reportlab', 'python-docx', 'PyPDF2', 'requests',
        'wikipedia', 'duckduckgo-search', 'python-dotenv'
    ]
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_requirements()

from colored import fg, bg, attr
from langchain.agents import AgentExecutor, create_react_agent
from langchain.tools import Tool
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import SystemMessage, HumanMessage
from langchain.prompts import PromptTemplate
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import PyPDF2
import wikipedia
from duckduckgo_search import DDGS

# =============================================================================
# CONFIGURATION VARIABLES
# =============================================================================
GEMINI_API_KEY = os.getenv('GOOGLE_API_KEY', 'your_gemini_api_key_here')
MAIN_MODEL = "gemini-2.0-flash-exp"  # Main LLM
SUB_MODEL = "gemini-2.0-flash-exp"   # Sub-agent tools
DATABASE_PATH = "data.db"
INPUT_DIR = Path("input")
OUTPUT_DIR = Path("output")

# Create directories
INPUT_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def colored_print(text: str, color: str = "white", style: str = "normal"):
    """Print colored text with formatting"""
    colors = {"red": fg.red, "green": fg.green, "blue": fg.blue, 
              "yellow": fg.yellow, "magenta": fg.magenta, "cyan": fg.cyan, "white": fg.white}
    styles = {"bold": attr.bold, "normal": attr.reset}
    print(f"{colors.get(color, fg.white)}{styles.get(style, attr.reset)}{text}{attr.reset}")

def encrypt_token(timestamp: str, name: str) -> str:
    """Generate encrypted token from timestamp and name"""
    data = f"{timestamp}_{name}"
    return hashlib.sha256(data.encode()).hexdigest()[:16]

def load_database() -> Dict:
    """Load user database from JSON file"""
    if not os.path.exists(DATABASE_PATH):
        return {}
    try:
        with open(DATABASE_PATH, 'r') as f:
            return json.load(f)
    except:
        return {}

def save_database(data: Dict):
    """Save user database to JSON file"""
    with open(DATABASE_PATH, 'w') as f:
        json.dump(data, f, indent=2)

def detect_location() -> str:
    """Detect user location via IP"""
    try:
        response = requests.get('http://ipapi.co/json/', timeout=5)
        data = response.json()
        return f"{data.get('city', 'Unknown')}, {data.get('country_name', 'Unknown')}"
    except:
        return "Location detection failed"

# =============================================================================
# TOOLS AND AGENTS
# =============================================================================

class LawLasTools:
    """Collection of legal assistant tools"""
    
    def __init__(self, llm):
        self.llm = llm
        self.ddgs = DDGS()
    
    def search_tool(self, query: str) -> str:
        """Search the web for legal information"""
        try:
            results = list(self.ddgs.text(query, max_results=3))
            return "\n".join([f"- {r['title']}: {r['body']}" for r in results])
        except Exception as e:
            return f"Search error: {str(e)}"
    
    def wikipedia_tool(self, topic: str) -> str:
        """Search Wikipedia for legal topics"""
        try:
            summary = wikipedia.summary(topic, sentences=3)
            return f"Wikipedia: {summary}"
        except wikipedia.exceptions.DisambiguationError as e:
            return f"Multiple topics found: {', '.join(e.options[:5])}"
        except Exception as e:
            return f"Wikipedia error: {str(e)}"
    
    def constitution_tool(self, country: str) -> str:
        """Fetch constitution information for a country"""
        query = f"{country} constitution text articles"
        try:
            results = list(self.ddgs.text(query, max_results=2))
            return f"Constitution of {country}:\n" + "\n".join([r['body'] for r in results])
        except Exception as e:
            return f"Constitution search error: {str(e)}"
    
    def file_analyzer(self, filename: str) -> str:
        """Analyze uploaded legal documents"""
        file_path = INPUT_DIR / filename
        if not file_path.exists():
            return f"File {filename} not found in /input directory"
        
        try:
            if filename.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return f"PDF Content:\n{text[:2000]}..."
            elif filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    return f"Text Content:\n{content[:2000]}..."
            else:
                return "Unsupported file format. Please use PDF or TXT files."
        except Exception as e:
            return f"File analysis error: {str(e)}"
    
    def document_generator(self, doc_type: str, content: str, filename: str) -> str:
        """Generate legal documents in PDF or DOCX format"""
        try:
            if filename.endswith('.pdf'):
                return self._generate_pdf(doc_type, content, filename)
            elif filename.endswith('.docx'):
                return self._generate_docx(doc_type, content, filename)
            else:
                return "Unsupported format. Use .pdf or .docx"
        except Exception as e:
            return f"Document generation error: {str(e)}"
    
    def _generate_pdf(self, doc_type: str, content: str, filename: str) -> str:
        """Generate PDF document"""
        output_path = OUTPUT_DIR / filename
        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, doc_type.upper())
        
        # Content
        c.setFont("Helvetica", 12)
        y_position = height - 100
        lines = content.split('\n')
        
        for line in lines:
            if y_position < 50:
                c.showPage()
                y_position = height - 50
            c.drawString(50, y_position, line[:80])  # Limit line length
            y_position -= 20
        
        c.save()
        return f"PDF generated successfully → /output/{filename}"
    
    def _generate_docx(self, doc_type: str, content: str, filename: str) -> str:
        """Generate DOCX document"""
        output_path = OUTPUT_DIR / filename
        doc = Document()
        
        # Title
        title = doc.add_heading(doc_type.upper(), 0)
        
        # Content
        doc.add_paragraph(content)
        
        # Footer
        footer_para = doc.add_paragraph()
        footer_para.add_run("Generated by LawLas ⚖️").italic = True
        
        doc.save(str(output_path))
        return f"DOCX generated successfully → /output/{filename}"
    
    def case_analyzer(self, case_text: str) -> str:
        """Analyze legal cases and provide applicable law insights"""
        prompt = f"""
        Analyze this legal case and provide insights on applicable law:
        
        Case: {case_text}
        
        Please provide:
        1. Key legal issues identified
        2. Applicable laws or precedents
        3. Potential outcomes
        4. Recommendations
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=prompt)])
            return f"Case Analysis:\n{response.content}"
        except Exception as e:
            return f"Case analysis error: {str(e)}"
    
    def legal_aid_finder(self, location: str) -> str:
        """Find budget-friendly legal aid in user's location"""
        query = f"free legal aid services {location} low income"
        try:
            results = list(self.ddgs.text(query, max_results=3))
            return "Legal Aid Resources:\n" + "\n".join([f"- {r['title']}: {r['body']}" for r in results])
        except Exception as e:
            return f"Legal aid search error: {str(e)}"
    
    def template_fetcher(self, doc_type: str) -> str:
        """Fetch document templates from online sources"""
        query = f"{doc_type} legal document template format structure"
        try:
            results = list(self.ddgs.text(query, max_results=2))
            return f"Template for {doc_type}:\n" + "\n".join([r['body'] for r in results])
        except Exception as e:
            return f"Template fetch error: {str(e)}"
    
    def risk_assessment(self, legal_issue: str) -> str:
        """Assess legal risks and provide escalation advice"""
        high_risk_keywords = ['criminal', 'felony', 'arrest', 'lawsuit', 'court', 'prison']
        
        risk_level = "LOW"
        if any(keyword in legal_issue.lower() for keyword in high_risk_keywords):
            risk_level = "HIGH"
        
        advice = "Consider consulting with a qualified attorney for personalized legal advice."
        if risk_level == "HIGH":
            advice = "⚠️  HIGH RISK: Strongly recommend immediate consultation with a qualified attorney."
        
        return f"Risk Assessment: {risk_level}\nAdvice: {advice}"
    
    def summarizer_tool(self, legal_text: str) -> str:
        """Convert complex legal text to plain language"""
        prompt = f"""
        Please explain this legal text in simple, plain English that a non-lawyer can understand:
        
        {legal_text}
        
        Use everyday language and provide examples where helpful.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=prompt)])
            return f"Plain English Summary:\n{response.content}"
        except Exception as e:
            return f"Summarization error: {str(e)}"

# =============================================================================
# MEMORY MANAGER
# =============================================================================

class MemoryManager:
    """Manage user conversation history"""
    
    def __init__(self, username: str):
        self.username = username
        self.db = load_database()
    
    def load_messages(self) -> List[Dict]:
        """Load user's conversation history"""
        if self.username in self.db:
            return self.db[self.username].get('messages', [])
        return []
    
    def save_message(self, role: str, content: str):
        """Save a message to user's history"""
        if self.username not in self.db:
            self.db[self.username] = {'name': self.username, 'messages': []}
        
        message = {
            'role': role,
            'content': content,
            'timestamp': datetime.now().isoformat()
        }
        
        self.db[self.username]['messages'].append(message)
        save_database(self.db)
    
    def get_recent_context(self, limit: int = 5) -> str:
        """Get recent conversation context"""
        messages = self.load_messages()
        recent = messages[-limit:] if len(messages) > limit else messages
        
        context = "Recent conversation:\n"
        for msg in recent:
            context += f"{msg['role']}: {msg['content'][:200]}...\n"
        
        return context

# =============================================================================
# MAIN LAWLAS SYSTEM
# =============================================================================

class LawLasSystem:
    """Main LawLas legal assistant system"""
    
    def __init__(self):
        self.current_user = None
        self.memory_manager = None
        
        # Initialize LLMs
        self.main_llm = ChatGoogleGenerativeAI(
            model=MAIN_MODEL,
            google_api_key=GEMINI_API_KEY,
            temperature=0.1
        )
        
        self.sub_llm = ChatGoogleGenerativeAI(
            model=SUB_MODEL,
            google_api_key=GEMINI_API_KEY,
            temperature=0.1
        )
        
        # Initialize tools
        self.tools_instance = LawLasTools(self.sub_llm)
        self.setup_tools()
    
    def setup_tools(self):
        """Setup LangChain tools"""
        self.tools = [
            Tool(
                name="SearchTool",
                description="Search the web for legal information",
                func=self.tools_instance.search_tool
            ),
            Tool(
                name="WikipediaTool", 
                description="Search Wikipedia for legal topics and explanations",
                func=self.tools_instance.wikipedia_tool
            ),
            Tool(
                name="ConstitutionTool",
                description="Fetch constitution information for any country",
                func=self.tools_instance.constitution_tool
            ),
            Tool(
                name="FileAnalyzer",
                description="Analyze uploaded legal documents (use filename)",
                func=self.tools_instance.file_analyzer
            ),
            Tool(
                name="DocumentGenerator",
                description="Generate legal documents in PDF or DOCX format",
                func=lambda x: self._parse_doc_generation(x)
            ),
            Tool(
                name="CaseAnalyzer",
                description="Analyze legal cases and provide applicable law insights",
                func=self.tools_instance.case_analyzer
            ),
            Tool(
                name="LegalAidFinder",
                description="Find budget-friendly legal aid services",
                func=self.tools_instance.legal_aid_finder
            ),
            Tool(
                name="TemplateFetcher",
                description="Fetch document templates for legal documents",
                func=self.tools_instance.template_fetcher
            ),
            Tool(
                name="RiskAssessment",
                description="Assess legal risks and provide escalation advice",
                func=self.tools_instance.risk_assessment
            ),
            Tool(
                name="Summarizer",
                description="Convert complex legal text to plain language",
                func=self.tools_instance.summarizer_tool
            ),
            Tool(
                name="LocationTool",
                description="Detect or get user location",
                func=lambda x: detect_location()
            )
        ]
    
    def _parse_doc_generation(self, input_text: str) -> str:
        """Parse document generation requests"""
        # Extract document type, content, and filename from input
        parts = input_text.split('--')
        if len(parts) < 2:
            return "Please specify filename with --filename.pdf format"
        
        content = parts[0].strip()
        filename = parts[1].strip()
        
        # Determine document type from content
        doc_type = "Legal Document"
        if "rent" in content.lower():
            doc_type = "Rental Agreement"
        elif "contract" in content.lower():
            doc_type = "Contract"
        elif "will" in content.lower():
            doc_type = "Last Will and Testament"
        
        return self.tools_instance.document_generator(doc_type, content, filename)
    
    def startup_menu(self):
        """Display startup menu and handle authentication"""
        colored_print("Welcome to LawLas ⚖️", "cyan", "bold")
        print()
        colored_print("1) Signup", "green")
        colored_print("2) Login", "blue")
        
        choice = input("\nSelect option (1-2): ").strip()
        
        if choice == "1":
            self.signup()
        elif choice == "2":
            self.login()
        else:
            colored_print("Invalid choice. Please try again.", "red")
            self.startup_menu()
    
    def signup(self):
        """User signup process"""
        name = input("Enter your name: ").strip()
        location = input("Enter your location (or press Enter to auto-detect): ").strip()
        
        if not location:
            location = detect_location()
        
        # Generate token
        timestamp = str(int(time.time()))
        token = encrypt_token(timestamp, name)
        
        # Save user data
        db = load_database()
        db[name] = {
            'name': name,
            'location': location,
            'token': token,
            'messages': []
        }
        save_database(db)
        
        colored_print(f"✓ Account created successfully! Welcome, {name}!", "green", "bold")
        self.current_user = name
        self.memory_manager = MemoryManager(name)
        self.chat_interface()
    
    def login(self):
        """User login process"""
        name = input("Enter your name: ").strip()
        token = input("Enter your token: ").strip()
        
        db = load_database()
        
        if name in db and db[name]['token'] == token:
            colored_print(f"✓ Login successful! Welcome back, {name}!", "green", "bold")
            self.current_user = name
            self.memory_manager = MemoryManager(name)
            
            # Load past messages
            messages = self.memory_manager.load_messages()
            if messages:
                colored_print(f"Loaded {len(messages)} previous messages.", "yellow")
            
            self.chat_interface()
        else:
            colored_print("✗ Invalid credentials. Please try again.", "red")
            self.startup_menu()
    
    def chat_interface(self):
        """Main chat interface"""
        colored_print(f"\n🤖 LawLas AI Legal Assistant", "magenta", "bold")
        colored_print("Type 'exit' to quit, 'help' for commands", "yellow")
        colored_print("⚠️  Disclaimer: This is for informational purposes only. Not legal advice.", "red")
        print()
        
        while True:
            try:
                user_input = input(f"{fg.cyan}You: {attr.reset}").strip()
                
                if user_input.lower() in ['exit', 'quit']:
                    colored_print("Goodbye! 👋", "cyan")
                    break
                
                if user_input.lower() == 'help':
                    self.show_help()
                    continue
                
                if not user_input:
                    continue
                
                # Save user message
                self.memory_manager.save_message('user', user_input)
                
                # Process user input
                response = self.process_user_input(user_input)
                
                # Display response with formatting
                colored_print(f"LawLas: {response}", "white")
                print()
                
                # Save AI response
                self.memory_manager.save_message('assistant', response)
                
            except KeyboardInterrupt:
                colored_print("\nGoodbye! 👋", "cyan")
                break
            except Exception as e:
                colored_print(f"Error: {str(e)}", "red")
    
    def process_user_input(self, user_input: str) -> str:
        """Process user input and generate response"""
        try:
            # Check for file operations
            if '--' in user_input:
                return self.handle_file_operation(user_input)
            
            # Get conversation context
            context = self.memory_manager.get_recent_context()
            
            # Create system prompt
            system_prompt = f"""
            You are LawLas, an AI legal assistant. You help users with legal questions and document analysis.
            
            User: {self.current_user}
            Location: {detect_location()}
            
            {context}
            
            Available tools: SearchTool, WikipediaTool, ConstitutionTool, FileAnalyzer, 
            DocumentGenerator, CaseAnalyzer, LegalAidFinder, TemplateFetcher, 
            RiskAssessment, Summarizer, LocationTool
            
            Always include appropriate disclaimers about legal advice.
            Be helpful, accurate, and conversational.
            """
            
            # Create agent prompt template
            template = """
            {system_prompt}
            
            Human: {input}
            
            Use available tools as needed to provide comprehensive legal assistance.
            
            {agent_scratchpad}
            """
            
            prompt = PromptTemplate(
                template=template,
                input_variables=["input", "agent_scratchpad", "system_prompt"]
            )
            
            # Create agent
            agent = create_react_agent(self.main_llm, self.tools, prompt)
            agent_executor = AgentExecutor(
                agent=agent, 
                tools=self.tools, 
                verbose=False,
                handle_parsing_errors=True
            )
            
            # Execute
            result = agent_executor.invoke({
                "input": user_input,
                "system_prompt": system_prompt
            })
            
            return self.format_response(result.get('output', 'No response generated.'))
            
        except Exception as e:
            return f"I apologize, but I encountered an error: {str(e)}. Please try rephrasing your question."
    
    def handle_file_operation(self, user_input: str) -> str:
        """Handle file operations (--filename syntax)"""
        if 'explain' in user_input.lower():
            # Extract filename
            parts = user_input.split('--')
            if len(parts) > 1:
                filename = parts[1].strip()
                return self.tools_instance.file_analyzer(filename)
        
        elif 'generate' in user_input.lower():
            # Handle document generation
            return self._parse_doc_generation(user_input)
        
        return "File operation not recognized. Use 'Explain --filename' or 'Generate ... --filename.pdf'"
    
    def format_response(self, response: str) -> str:
        """Format AI response with colors and styling"""
        # Add disclaimer
        if 'disclaimer' not in response.lower():
            response += "\n\n⚠️ Disclaimer: This information is for educational purposes only and does not constitute legal advice."
        
        # Bold important terms (simple implementation)
        legal_terms = ['contract', 'liability', 'negligence', 'breach', 'damages', 'rights', 'obligations']
        for term in legal_terms:
            response = response.replace(term, f"{attr.bold}{term}{attr.reset}")
        
        return response
    
    def show_help(self):
        """Display help information"""
        colored_print("\n📋 LawLas Commands:", "cyan", "bold")
        print("• Ask any legal question in natural language")
        print("• 'Explain --filename.txt' - Analyze uploaded documents")
        print("• 'Generate rental agreement --rent.pdf' - Create documents")
        print("• 'Find legal aid in [location]' - Budget legal services")
        print("• 'Analyze case: [case text]' - Legal case analysis")
        print("• 'Constitution of [country]' - Constitutional information")
        print("• 'exit' or 'quit' - Exit the program")
        colored_print("\n💡 Upload files to /input directory for analysis", "yellow")
        colored_print("📁 Generated documents saved to /output directory", "yellow")
        print()

# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main execution function"""
    try:
        # Initialize system
        lawlas = LawLasSystem()
        
        # Start with authentication menu
        lawlas.startup_menu()
        
    except KeyboardInterrupt:
        colored_print("\nProgram interrupted. Goodbye! 👋", "cyan")
    except Exception as e:
        colored_print(f"Fatal error: {str(e)}", "red")
        colored_print("Please check your configuration and try again.", "yellow")

if __name__ == "__main__":
    main()